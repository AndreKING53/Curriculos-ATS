from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

# Acessa um estilo específico do documento
def configurar_estilo(doc):
    normal_style = doc.styles['Normal']
    font = normal_style.font
# Configura o estilo da fonte para Times New Roman, tamanho 12, sem negrito, sem itálico e cor preta
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.bold = False
    font.italic = False

def titulo_estilo(doc):
    try:
        titulo_style = doc.styles['Titulo']
    except KeyError:
        titulo_style = doc.styles.add_style('Titulo', 1)
        titulo_font = titulo_style.font
        titulo_font.name = 'Times New Roman'
        titulo_font.size = Pt(16)
    titulo_font.bold = True

def negrito(doc):
    negrito_style = doc.styles.add_style('Negrito', 1)
    negrito_font = negrito_style.font
    negrito_font.name = 'Times New Roman'
    negrito_font.size = Pt(12)
    negrito_font.bold = True

def adicionar_titulo(doc, dados):
    paragrafo_nome = doc.add_paragraph(dados["nome"], style="Titulo")
    paragrafo_nome.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(dados["localizacao"], style="Negrito")

def adicionar_contato(doc, dados):
    contato = dados["contato"]
    linha = " | ".join(
        valor
        for valor in contato.values()
        if valor
    )
    doc.add_paragraph(linha)

    paragrafo_objetivo = doc.add_paragraph()
    paragrafo_objetivo.add_run('Objetivo: ').bold = True
    paragrafo_objetivo.add_run(dados["objetivo"])


def adicionar_formacao(doc, dados):
    doc.add_paragraph("Formação Acadêmica", style='Titulo')
    for formacao in dados["formaçao"]:
        doc.add_paragraph(
        f'{formacao["instituicao"]} - {formacao["curso"]} ({formacao["inicio"]} - {formacao["fim"]})'
    )

def adicionar_experiencias(doc, dados):
    doc.add_paragraph("Experiências Profissionais", style='Titulo')
    for experiencia in dados["experiencias"]:
        doc.add_paragraph(
        f'{experiencia["empresa"]} - {experiencia["cargo"]} ({experiencia["data"]})'
        )
        doc.add_paragraph(experiencia["descricao"], style="List Bullet")

def adicionar_competencias(doc, dados):
    doc.add_paragraph("Competências", style='Titulo')
    for competencia in dados["competencias"]:
        doc.add_paragraph(competencia, style="List Bullet")

def adicionar_idiomas(doc, dados):
    doc.add_paragraph("Idiomas")
    for idioma in dados["idiomas"]:
        doc.add_paragraph(
        f'{idioma["idioma"]} - {idioma["nivel"]}',
        style="List Bullet"
    )

COR_PADRAO = RGBColor(0, 0, 0)